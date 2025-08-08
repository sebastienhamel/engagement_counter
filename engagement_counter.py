from docx import Document
from collections import defaultdict
import os
import pandas as pd
import re
import datetime

class EngagementCount():
    """
    A class to handle engagement counting from Word documents.
    It extracts interventions by speakers and aggregates them by date.
    """

    def __init__(self, folder_path):
        self.folder_path = folder_path
        self.instructor_name = self.get_instructor_name()
    
    def extract_files_from_folder(self, folder_path, file_pattern=""):
        """
        Extracts all relevant files from the specified folder.
        Returns a list of filenames.
        """
        files = []
        
        for filename in os.listdir(folder_path):
            if filename.startswith(file_pattern) and filename.endswith(".docx"):
                files.append(filename)
        
        return files

    def perform_file_date_mapping(self, files):
        """
        Maps filenames to their respective dates.
        """
        file_date_mapping = {}
        
        for filename in files:
            full_path = os.path.join(folder_path, filename)
            date = self.get_date_from_file(full_path)
            
            if date:
                print(f"Extracted date {date} from {filename}")
                file_date_mapping[filename] = date
            else:
                raise ValueError(f"Could not extract date from file: {filename}")

        if not files:
            print("No relevant files found in the specified folder.")
        
        return file_date_mapping
        
    def extract_interventions(self, doc_path):
        """Extracts interventions from a Word document.
        Returns a dictionary with speaker names as keys and their intervention counts as values.
        """
        document = Document(doc_path)
        text = "\n".join([para.text for para in document.paragraphs])
        speaker_pattern = re.compile(r"^([A-Z][a-z]+(?: [A-Z][a-z\-]+)*?)\s+", re.MULTILINE)
        speakers = speaker_pattern.findall(text)
        counts = defaultdict(int)
        
        for speaker in speakers:
            counts[speaker.strip()] += 1
        
        return counts

    def get_date_from_file(self, filepath):
        """
        Extracts the date from the file content or filename.
        Handles formats like:
        - Week 4 - Day 2 - AM Session-20250708_093029-Meeting Recording
        - July 8, 2025, 8:30AM
        - 2025-07-08
        """
        
        # Try to extract YYYY-MM-DD from filename first
        match = re.search(r"(\d{4}-\d{2}-\d{2})", filepath)
        if match:
            return match.group(1)

        # Try to extract YYYYMMDD from filename
        match = re.search(r"(\d{8})_", filepath)
        if match:
            try:
                return datetime.datetime.strptime(match.group(1), "%Y%m%d").strftime("%Y-%m-%d")
            except ValueError:
                pass

        # Try to extract from file content
        try:
            document = Document(filepath)
            
            for para in document.paragraphs[:3]:  # Only check first few lines
                text = para.text.strip()
                
                # Try "Month Day, Year" format
                match = re.search(r"([A-Za-z]+ \d{1,2}, \d{4})", text)
                if match:
                    try:
                        return datetime.datetime.strptime(match.group(1), "%B %d, %Y").strftime("%Y-%m-%d")
                    except ValueError:
                        pass
                
                # Try YYYY-MM-DD
                match = re.search(r"(\d{4}-\d{2}-\d{2})", text)
                if match:
                    return match.group(1)
                
                # Try YYYYMMDD
                match = re.search(r"(\d{8})", text)
                if match:
                    try:
                        return datetime.datetime.strptime(match.group(1), "%Y%m%d").strftime("%Y-%m-%d")
                    except ValueError:
                        pass
        
        except Exception as e:
            print(f"Error reading file for date extraction: {filepath}: {e}")
        
        return None 

    def extract_interventions_from_files(self, file_date_mapping):
        """
        Extracts interventions from files and aggregates them by date and speaker.
        """
        summary = defaultdict(lambda: defaultdict(int))
        
        for filename, date in file_date_mapping.items():
            full_path = os.path.join(folder_path, filename)
            
            try:
                if not os.path.exists(full_path):
                    print(f"File not found: {full_path}")
                    continue
                daily_counts = self.extract_interventions(full_path)
            except Exception as e:
                print(f"Error processing file {full_path}: {e}")
            
            for speaker, count in daily_counts.items():
                summary[date][speaker] += count
        
        return summary

    def flatten_summary_to_dataframe(self, summary):
        """
        Flattens the summary dictionary into a DataFrame, excluding the instructor's interventions.
        """
        data = []
        
        for date, speakers in summary.items():
            for speaker, count in speakers.items():
                if self.instructor_name and speaker == self.instructor_name:
                    print(f"Skipping instructor {self.instructor_name} for date {date}")
                    continue
                
                data.append({"Date": date, "Speaker": speaker, "Intervention Count": count})
        
        return pd.DataFrame(data).sort_values(by=["Date", "Speaker"])

    def save_summary_to_csv(self, df, folder_path):
        """
        Saves the DataFrame to a CSV file in the specified folder.
        """
        if not os.path.exists(folder_path):
            os.makedirs(folder_path)

        output_file = os.path.join(folder_path, "engagement_summary.csv")
        df.to_csv(output_file, index=False)
        print(f"Engagement summary saved to {output_file}")

    def get_instructor_name(self):
        """
        Returns the instructor's name.
        """
        instructor_name_file = "instructor_name.txt"
        with open(instructor_name_file, "r", encoding="utf-8") as f:
            instructor_name = f.read().strip()
        return instructor_name or None


if __name__ == "__main__":
    folder_path = "transcripts/" # Replace with the actual folder path if needed
    # Read instructor name from a separate file (e.g., instructor_name.txt)
    file_pattern = "Week "  # Pattern to match relevant files
    engagement_counter = EngagementCount(folder_path)
    
    print("Starting engagement counter script...")
    files = engagement_counter.extract_files_from_folder(folder_path, file_pattern)
    file_date_mapping = engagement_counter.perform_file_date_mapping(files)
    summary = engagement_counter.extract_interventions_from_files(file_date_mapping)
    df = engagement_counter.flatten_summary_to_dataframe(summary)
    engagement_counter.save_summary_to_csv(df, folder_path)
    print("Engagement counter script completed successfully.")



